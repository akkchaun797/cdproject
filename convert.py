from docx import Document
from docx.shared import Pt, Inches
import re

with open(r'C:\Users\AKKU\OneDrive\文档\GitHub\cdproject\plan.md', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1

    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1

    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1

    elif line.strip() in ('---', '***', '___'):
        doc.add_paragraph('')
        i += 1

    elif line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        i += 1
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|---'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            table_lines.append(row)
            i += 1
        data_rows = [r for r in table_lines if not all(set(c) <= set('-: ') for c in r)]
        if data_rows:
            cols = max(len(r) for r in data_rows)
            t = doc.add_table(rows=len(data_rows), cols=cols)
            t.style = 'Light List Accent 1'
            for ri, row in enumerate(data_rows):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ''
                    t.rows[ri].cells[ci].text = strip_md(cell_text)

    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(strip_md(line[2:].strip()), style='List Bullet')
        i += 1

    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line).strip()
        doc.add_paragraph(strip_md(text), style='List Number')
        i += 1

    elif line.startswith('> '):
        text = strip_md(line[2:].strip())
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.4)
        if p.runs:
            p.runs[0].italic = True

    elif line.strip() == '':
        i += 1

    else:
        doc.add_paragraph(strip_md(line))
        i += 1

doc.save(r'C:\Users\AKKU\OneDrive\文档\GitHub\cdproject\plan.docx')
print('plan.docx created successfully.')
